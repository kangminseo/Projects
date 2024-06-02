import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 정렬 설정을 위해 추가

def load_document():
    global doc
    filepath = filedialog.askopenfilename()
    doc = Document(filepath)
    status_var.set(f"Loaded document: {filepath}")

def save_document():
    global doc
    filepath = filedialog.asksaveasfilename(defaultextension=".docx")
    doc.save(filepath)
    status_var.set(f"Saved modified document: {filepath}")

def update_document():
    global doc
    # 문서 텍스트 수정
    text_modifications = {
        "수정부제목": subtitle_entry.get(),
    }

    for para in doc.paragraphs:
        for key, value in text_modifications.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

        # 첫 번째 표 수정 - 수정수신자와 수정제목
        if doc.tables:
            table1 = doc.tables[0]  # 첫 번째 표
            # '수정수신자'와 '수정제목' 단어가 포함된 셀 내용만 교체
            for i, row in enumerate(table1.rows):
                for cell in row.cells:
                    if "수정교수님" in cell.text:
                        cell.text = cell.text.replace("수정교수님", receiver_entry.get())
                    if "수정제목" in cell.text:
                        cell.text = cell.text.replace("수정제목", title_entry.get())


    # 두 번째 표 수정 - 나머지 수정사항
    if len(doc.tables) > 1:
        table2 = doc.tables[1]  # 두 번째 표
        modifications = [
            (date_entry.get(), '수정일시'),
            (subject_entry.get(), '수정과목'),
            (section_entry.get(), '수정분반'),
            (department_entry.get(), '수정학과'),
            (student_id_entry.get(), '수정학번'),
            (name_entry.get(), '수정이름'),
        ]

        for col_idx, (new_text, _) in enumerate(modifications):
            cell = table2.cell(1, col_idx)
            cell.text = new_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 세 번째 표 수정 - '수정시행번호'와 '수정날짜'
        if len(doc.tables) > 2:
            table3 = doc.tables[2]  # 세 번째 표
            execution_number_row = table3.rows[3]  # 네 번째 행
            for cell in execution_number_row.cells:
                if "수정시행번호" in cell.text:
                    # '수정시행번호' 단어만 교체
                    cell.text = cell.text.replace("수정시행번호", execution_no_entry.get())
                if "수정날짜" in cell.text:
                    # '수정날짜' 단어만 교체
                    cell.text = cell.text.replace("수정날짜", modification_date_entry.get())

        status_var.set("Document updated successfully.")

# GUI setup
root = tk.Tk()
root.title("37대총동부회장 강민서 제작")

# Entries for modifications
labels = ["교수님성함", "제목", "협조문", "일시", "과목", "분반", "학과", "학번", "이름", "시행번호", "날짜"]
entries = []

for i, label in enumerate(labels):
    tk.Label(root, text=label).grid(row=i, column=0)
    entry = tk.Entry(root)
    entry.grid(row=i, column=1)
    entries.append(entry)

(receiver_entry, title_entry, subtitle_entry, date_entry, subject_entry, section_entry,
 department_entry, student_id_entry, name_entry, execution_no_entry,
 modification_date_entry) = entries

# Buttons
tk.Button(root, text="파일 불러오기", command=load_document).grid(row=11, column=0)
tk.Button(root, text="업데이트", command=update_document).grid(row=11, column=1)
tk.Button(root, text="업뎃파일저장", command=save_document).grid(row=11, column=2)

# Status bar
status_var = tk.StringVar()
status_var.set("Ready")
status_bar = tk.Label(root, textvariable=status_var, bd=1, relief=tk.SUNKEN, anchor=tk.W)
status_bar.grid(row=12, column=0, columnspan=3, sticky=tk.W + tk.E)

root.mainloop()